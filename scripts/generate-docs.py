import sys
import os

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x10, 0x18, 0x20)
        else:
            run.font.color.rgb = RGBColor(0x1C, 0x2A, 0x3D)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

OUTPUT_DIR = '/home/z/my-project/download'

# ============================================================
# DOC 1: TECHNICAL DOC (FOR DEV)
# ============================================================

def create_dev_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.3

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('FleetControl')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x10, 0x18, 0x20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Documentacao Tecnica do Sistema')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x50, 0x60, 0x70)
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Ride Hailing Corporativo | Backend + Dashboard Administrativo')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x90, 0xA0)
    doc.add_paragraph()
    doc.add_page_break()

    # 1. VISAO GERAL
    add_heading_styled(doc, '1. Visao Geral da Arquitetura', 1)
    add_body(doc, 'O FleetControl e um sistema completo de gerenciamento de frota e viagens corporativas, construido com Next.js 16 (App Router), Prisma ORM, TypeScript e Tailwind CSS 4. A aplicacao combina um backend API robusto com um dashboard administrativo single-page dentro do mesmo projeto Next.js, alem de um microservico WebSocket independente para rastreamento em tempo real.')
    add_body(doc, 'A arquitetura segue o modelo monorepo simplificado: o Next.js serve tanto as API routes (/api/*) quanto as paginas do dashboard. O WebSocket de rastreamento opera como um servico independente na porta 3003. O banco de dados atualmente utiliza SQLite para desenvolvimento, com schema Prisma pronto para migracao para PostgreSQL em producao.')

    add_heading_styled(doc, '1.1 Stack Tecnologico', 2)
    add_table(doc,
        ['Camada', 'Tecnologia', 'Versao', 'Observacao'],
        [
            ['Framework', 'Next.js (App Router)', '16.x', 'API Routes + SSR/CSR'],
            ['Linguagem', 'TypeScript', '5.x', 'Strict mode'],
            ['Estilo', 'Tailwind CSS', '4.x', 'Com shadcn/ui'],
            ['ORM', 'Prisma', '6.x', 'SQLite (dev) / PostgreSQL (prod)'],
            ['Autenticacao', 'JWT + bcryptjs', '-', 'Access + Refresh tokens'],
            ['WebSocket', 'Socket.io', '4.x', 'Microservico independente'],
            ['Relatorios', 'xlsx (SheetJS)', '0.18.x', 'Exportacao XLSX/CSV'],
            ['Estado Cliente', 'Zustand', '5.x', 'Auth store'],
            ['Runtime', 'Bun', 'latest', 'Desenvolvimento e build'],
        ]
    )

    # 2. ESTRUTURA DE DIRETORIOS
    add_heading_styled(doc, '2. Estrutura de Diretorios', 1)
    add_code_block(doc, '''project-root/
+-- prisma/
|   +-- schema.prisma       # Schema completo (12 entidades)
|   +-- seed.ts             # Seed com dados de demo
+-- src/
|   +-- app/
|   |   +-- page.tsx        # Entry point (DashboardShell)
|   |   +-- layout.tsx      # Layout raiz
|   |   +-- globals.css     # Estilos + tema
|   |   +-- api/
|   |       +-- auth/       # login, logout, refresh, me
|   |       +-- users/      # CRUD usuarios
|   |       +-- vehicles/   # CRUD veiculos + metadata
|   |       +-- drivers/    # CRUD motoristas
|   |       +-- passengers/ # CRUD passageiros
|   |       +-- cost-centers/ # CRUD + metadata
|   |       +-- rides/      # CRUD + dispatch + status
|   |       +-- rules/      # Regras disponibilidade
|   |       +-- checkouts/  # Check-in/Check-out
|   |       +-- reports/    # Exportacao XLSX/CSV
|   |       +-- audit-logs/ # Logs auditoria
|   |       +-- health/     # Health check
|   |       +-- metrics/    # Metricas
|   +-- lib/
|   |   +-- auth.ts         # JWT, hash, rate limiting
|   |   +-- auth-middleware.ts  # Validacao token
|   |   +-- audit.ts        # Logger auditoria
|   |   +-- state-machine.ts # Transicoes estado
|   |   +-- geofencing.ts   # Haversine, horarios
|   |   +-- db.ts           # Prisma Client
|   |   +-- api.ts          # Client HTTP com refresh
|   +-- types/index.ts      # Tipos e interfaces
|   +-- stores/auth-store.ts # Zustand (auth)
|   +-- components/
|       +-- dashboard/      # Login + Shell (11 secoes)
|       +-- ui/              # shadcn/ui
+-- mini-services/tracking-service/  # WebSocket (porta 3003)
+-- vercel.json
+-- .env.example''')

    # 3. MODELO DE DADOS
    add_heading_styled(doc, '3. Modelo de Dados (Prisma Schema)', 1)
    add_body(doc, 'O schema define 12 entidades com relacoes completas:')
    add_table(doc,
        ['Entidade', 'Descricao', 'Campos Principais'],
        [
            ['User', 'Usuarios do sistema', 'email, name, passwordHash, role, branchId, branchName, isActive'],
            ['Vehicle', 'Veiculos da frota', 'plate (unique), model, capacity, trackerId, status, color, year'],
            ['VehicleMetadata', 'Metadados KV do veiculo', 'vehicleId, key, value'],
            ['Driver', 'Motoristas', 'userId (unique), licenseNumber, licenseExpiry, phone, status, currentVehicleId'],
            ['Passenger', 'Passageiros', 'userId (unique), phone, costCenterId'],
            ['CostCenter', 'Centros de custo', 'name, code (unique), description, isActive'],
            ['CostCenterMetadata', 'Metadados KV do CC', 'costCenterId, key, value'],
            ['AvailabilityRule', 'Regras disponibilidade', 'name, centerLat, centerLng, radiusKm, allowedDays, startTime, endTime'],
            ['Ride', 'Viagens', 'passengerId, driverId, vehicleId, costCenterId, status, pickup/dropoff coords+address, timestamps'],
            ['VehicleCheckout', 'Check-in/Check-out', 'vehicleId, driverId, checkedOutAt, checkedInAt, mileageOut/In, fuelLevelOut/In, status'],
            ['AuditLog', 'Logs auditoria (append-only)', 'userId, action, resource, resourceId, details (JSON), ipAddress, userAgent'],
            ['RefreshToken', 'Tokens de refresh', 'token (unique), userId, expiresAt'],
        ]
    )

    # 4. AUTENTICACAO
    add_heading_styled(doc, '4. Sistema de Autenticacao', 1)
    add_heading_styled(doc, '4.1 JWT Access + Refresh Tokens', 2)
    add_body(doc, 'O login retorna um accessToken (JWT, 15min) no body e um refreshToken (UUID, 7 dias) como cookie httpOnly. O client (src/lib/api.ts) armazena o accessToken no localStorage. Quando o accessToken expira, tenta automaticamente o refresh via /api/auth/refresh usando o cookie. Se tambem expirou, redireciona para login.')

    add_heading_styled(doc, '4.2 RBAC (Role-Based Access Control)', 2)
    add_table(doc,
        ['Role', 'Descricao', 'Permissoes'],
        [
            ['SUPER_ADMIN', 'Administrador global', 'Acesso total. CRUD em tudo. Auditoria e Rules.'],
            ['MANAGER', 'Gerente de filial', 'CRUD usuarios, veiculos, motoristas, passageiros, CC, viagens, checkouts.'],
            ['DRIVER', 'Motorista', 'Visualiza/altera status das proprias viagens.'],
            ['PASSENGER', 'Passageiro', 'Cria solicitacoes e visualiza proprias viagens.'],
        ]
    )

    add_heading_styled(doc, '4.3 Rate Limiting', 2)
    add_body(doc, 'Bloqueia apos 5 tentativas falhas por email em 15 minutos (em memoria, Map). Para producao, recomenda-se Redis.')

    # 5. MAQUINA DE ESTADOS
    add_heading_styled(doc, '5. Maquina de Estados', 1)
    add_heading_styled(doc, '5.1 Veiculo', 2)
    add_table(doc, ['Estado', 'Transicoes'], [
        ['AVAILABLE', 'EN_ROUTE, OFFLINE, MAINTENANCE'],
        ['EN_ROUTE', 'IN_RIDE, AVAILABLE'],
        ['IN_RIDE', 'AVAILABLE'],
        ['OFFLINE', 'AVAILABLE'],
        ['MAINTENANCE', 'AVAILABLE'],
    ])
    add_heading_styled(doc, '5.2 Viagem', 2)
    add_table(doc, ['Estado', 'Transicoes'], [
        ['REQUESTED', 'DISPATCHED, CANCELED'],
        ['DISPATCHED', 'ARRIVED_AT_PICKUP, CANCELED'],
        ['ARRIVED_AT_PICKUP', 'IN_PROGRESS, CANCELED'],
        ['IN_PROGRESS', 'COMPLETED, CANCELED'],
        ['COMPLETED', '(final)'],
        ['CANCELED', '(final)'],
    ])

    # 6. API REFERENCE
    add_heading_styled(doc, '6. API Routes - Referencia Completa', 1)
    add_table(doc,
        ['Metodo', 'Endpoint', 'Auth', 'Descricao'],
        [
            ['POST', '/api/auth/login', 'Publico', 'Login. Retorna accessToken + refreshToken.'],
            ['POST', '/api/auth/refresh', 'Cookie', 'Renova accessToken.'],
            ['GET', '/api/auth/me', 'JWT', 'Dados do usuario autenticado.'],
            ['POST', '/api/auth/logout', 'JWT', 'Limpa refreshToken.'],
            ['GET', '/api/users', 'SA/MG', 'Lista com paginacao (page, limit, role, search).'],
            ['POST', '/api/users', 'SA', 'Cria usuario.'],
            ['GET|PUT|DEL', '/api/users/[id]', 'SA/MG', 'CRUD usuario.'],
            ['GET', '/api/vehicles', 'SA/MG', 'Lista veiculos (status, search).'],
            ['POST', '/api/vehicles', 'SA/MG', 'Cria veiculo + metadata[].'],
            ['GET|PUT|DEL', '/api/vehicles/[id]', 'SA/MG', 'CRUD. PUT valida state machine.'],
            ['GET|POST|DEL', '/api/vehicles/[id]/metadata', 'SA/MG', 'Metadata KV.'],
            ['GET|POST', '/api/drivers', 'SA/MG', 'Lista/Cria motoristas.'],
            ['GET|PUT|DEL', '/api/drivers/[id]', 'SA/MG', 'CRUD motorista.'],
            ['GET|POST', '/api/passengers', 'SA/MG', 'Lista/Cria passageiros.'],
            ['GET|PUT|DEL', '/api/passengers/[id]', 'SA/MG', 'CRUD passageiro.'],
            ['GET|POST', '/api/cost-centers', 'SA/MG', 'Lista/Cria centros de custo.'],
            ['GET|PUT|DEL', '/api/cost-centers/[id]', 'SA/MG', 'CRUD.'],
            ['GET|POST', '/api/cost-centers/[id]/metadata', 'SA/MG', 'Metadata KV.'],
            ['GET|POST', '/api/rules', 'SA', 'Regras de disponibilidade.'],
            ['GET|PUT|DEL', '/api/rules/[id]', 'SA', 'CRUD regras.'],
            ['GET', '/api/rides', 'SA/MG', 'Lista viagens com filtros.'],
            ['POST', '/api/rides', 'Auth', 'Cria viagem (valida geofencing).'],
            ['GET|PUT', '/api/rides/[id]', 'Auth', 'Detalhe / Atualiza status.'],
            ['POST', '/api/rides/[id]/dispatch', 'SA/MG', 'Despacha viagem (driver+vehicle).'],
            ['GET|POST', '/api/checkouts', 'SA/MG', 'Check-out de veiculo.'],
            ['POST', '/api/checkouts/[id]/return', 'SA/MG', 'Retorna veiculo.'],
            ['GET', '/api/reports/rides', 'SA/MG', 'Exporta (json/xlsx/csv).'],
            ['GET', '/api/audit-logs', 'SA', 'Logs auditoria.'],
            ['GET', '/api/health', 'Publico', 'Health check.'],
            ['GET', '/api/metrics', 'Auth', 'Metricas do sistema.'],
        ]
    )
    add_body(doc, 'SA = SUPER_ADMIN, MG = MANAGER. Response format: { success: boolean, data?: T, error?: string, pagination?: { page, limit, total, totalPages } }')

    # 7. GEOFENCING
    add_heading_styled(doc, '7. Geofencing (src/lib/geofencing.ts)', 1)
    add_body(doc, 'isWithinRadius(lat1, lng1, lat2, lng2, radiusKm) - Haversine. isWithinAllowedTime(startTime, endTime). isAllowedDay(allowedDays). validateRideRequest(pickupLat, pickupLng, rules[]) - Valida contra todas as regras ativas.')

    # 8. WEBSOCKET
    add_heading_styled(doc, '8. Servico WebSocket (porta 3003)', 1)
    add_table(doc, ['Evento', 'Tipo', 'Descricao'], [
        ['join-dashboard', 'in', 'Canal global da frota.'],
        ['leave-dashboard', 'in', 'Sai do canal.'],
        ['join-ride', 'in', 'Canal de viagem (rideId).'],
        ['leave-ride', 'in', 'Sai do canal.'],
        ['simulate-tracking', 'in', 'Inicia GPS simulado (demo).'],
        ['stop-tracking', 'in', 'Para simulacao.'],
        ['vehicle-location', 'out', '{vehicleId, lat, lng, heading, speed, timestamp}.'],
    ])

    # 9. AUDIT
    add_heading_styled(doc, '9. Auditoria (src/lib/audit.ts)', 1)
    add_body(doc, 'Toda mutacao gera AuditLog (append-only). Registra: userId, action, resource, resourceId, details (JSON diff), ipAddress, userAgent.')

    # 10. DASHBOARD
    add_heading_styled(doc, '10. Dashboard (src/components/dashboard/)', 1)
    add_table(doc, ['Secao', 'Funcionalidades'], [
        ['Painel Geral', 'Cards metricas, distribuicao status, status sistema.'],
        ['Usuarios', 'CRUD com busca e paginacao.'],
        ['Veiculos', 'CRUD com metadados + state machine.'],
        ['Motoristas', 'CRUD vinculando usuario.'],
        ['Passageiros', 'CRUD com centro de custo.'],
        ['Centros de Custo', 'CRUD com metadados.'],
        ['Viagens', 'Filtro status, dispatch, alteracao status.'],
        ['Regras', 'CRUD geofencing + horarios.'],
        ['Frota (Check)', 'Check-out/check-in veiculos.'],
        ['Relatorios', 'Exportacao XLSX/CSV/JSON.'],
        ['Auditoria', 'Logs com filtros e paginacao.'],
    ])

    # 11. DEPLOY
    add_heading_styled(doc, '11. Deploy', 1)
    add_heading_styled(doc, '11.1 Vercel', 2)
    add_code_block(doc, '''1. Conecte o repositorio no Vercel
2. Environment Variables:
   DATABASE_URL=<PostgreSQL connection string>
   JWT_SECRET=<chave forte>
   JWT_REFRESH_SECRET=<outra chave forte>
3. Build detectado automaticamente (vercel.json)''')
    add_heading_styled(doc, '11.2 Comandos Locais', 2)
    add_code_block(doc, '''bun install
bun run db:push          # Gerar schema no banco
bun run prisma/seed.ts   # Dados de demo
bun run dev             # Desenvolvimento (porta 3000)
cd mini-services/tracking-service && bun install && bun run dev  # WebSocket (3003)''')

    # 12. CONTAS
    add_heading_styled(doc, '12. Contas de Demo (Seed)', 1)
    add_table(doc, ['Perfil', 'Email', 'Senha'], [
        ['Super Admin', 'admin@corporate.com', 'Admin@123'],
        ['Gerente', 'manager@corporate.com', 'Manager@123'],
        ['Motorista', 'driver1@corporate.com', 'Driver@123'],
        ['Passageiro', 'passenger1@corporate.com', 'Passenger@123'],
    ])
    add_body(doc, 'Seed cria: 3 veiculos, 2 centros de custo (CC-001 Marketing, CC-002 Financeiro), 1 regra (dias uteis 8h-18h, 50km SP), 1 viagem concluida.')

    # 13. PROXIMOS PASSOS
    add_heading_styled(doc, '13. Recomendacoes Producao', 1)
    add_body(doc, 'Migrar SQLite para PostgreSQL. Rate limiting com Redis. Integrar API GPS real no WebSocket. Swagger/OpenAPI. Testes automatizados. CORS restritivo. Notificacoes push/email. Prisma Migrate para versionamento do schema.')

    path = os.path.join(OUTPUT_DIR, 'FleetControl-Documentacao-Tecnica.docx')
    doc.save(path)
    print(f'Saved: {path}')


# ============================================================
# DOC 2: CLIENT DOC (GENERAL)
# ============================================================

def create_client_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.3

    # COVER
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('FleetControl')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x10, 0x18, 0x20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Plataforma de Gestao de Frota Corporativa')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x50, 0x60, 0x70)
    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('Documento de Entrega Tecnica')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x90, 0xA0)
    doc.add_page_break()

    # 1
    add_heading_styled(doc, '1. Introducao', 1)
    add_body(doc, 'O FleetControl e uma plataforma web completa para gestao de operacoes de transporte corporativo. O sistema permite controlar toda a cadeia operacional de viagens: desde a solicitacao pelo passageiro ate o checkout do veiculo pelo motorista, incluindo rastreamento em tempo real, controle de acesso por perfis, regras de disponibilidade geograficas e horarias, alocacao por centros de custo, e emissao de relatorios operacionais.')
    add_body(doc, 'A solucao foi desenvolvida com tecnologias modernas e escalaveis, utilizando Next.js como framework principal, o que garante performance e facilidade de deploy em plataformas cloud como Vercel. A arquitetura foi projetada para suportar crescente volume de dados e pode ser integrada com APIs de rastreamento GPS existentes.')

    # 2
    add_heading_styled(doc, '2. Funcionalidades Implementadas', 1)

    add_heading_styled(doc, '2.1 Gestao de Usuarios e Acesso', 2)
    add_body(doc, 'O sistema possui controle de acesso baseado em perfis (RBAC), com quatro niveis hierarquicos que definem exatamente o que cada usuario pode visualizar e operar. Cada usuario pode ser vinculado a uma filial corporativa, permitindo que gestores vejam apenas os dados da sua unidade. A autenticacao utiliza tokens JWT com renovacao automatica, e ha protecao contra tentativas de acesso indevido. As senhas sao armazenadas com criptografia unidirecional.')

    add_heading_styled(doc, '2.2 Gestao de Frota', 2)
    add_body(doc, 'Cadastro completo de veiculos com dados como placa, modelo, capacidade, cor, ano e identificador do rastreador GPS. Cada veiculo possui um ciclo de vida controlado por estados (Disponivel, A Caminho, Em Viagem, Offline, Manutencao), e o sistema impede transicoes invalidas automaticamente. Metadados customizados podem ser anexados para informacoes adicionais conforme necessidade da operacao.')

    add_heading_styled(doc, '2.3 Gestao de Motoristas e Passageiros', 2)
    add_body(doc, 'Motoristas sao cadastrados com dados profissionais como numero e validade da CNH, telefone e vinculo com o usuario do sistema. Passageiros sao vinculados a centros de custo para controle de rateio de despesas. Ambos os perfis herdam as permissoes de acesso do usuario ao qual estao vinculados.')

    add_heading_styled(doc, '2.4 Centros de Custo', 2)
    add_body(doc, 'Centros de custo permitem alocar despesas de viagens a diferentes departamentos ou projetos da empresa. Cada viagem pode ser vinculada a um centro de custo, possibilitando relatorios de rateio e controle orcamentario. Assim como veiculos, centros de custo suportam metadados customizados.')

    add_heading_styled(doc, '2.5 Operacoes de Viagem', 2)
    add_body(doc, 'O modulo de viagens implementa o fluxo completo: o passageiro solicita uma viagem informando origem e destino; o gestor despacha a viagem atribuindo um motorista e veiculo disponivel; o motorista acompanha o percurso; e a viagem e finalizada. Cada etapa e registrada com timestamp, e o sistema valida automaticamente se a solicitacao esta dentro das regras de disponibilidade (area geografica, dias e horarios permitidos).')
    add_body(doc, 'Os estados de uma viagem sao: Solicitada, Despachada, No Local de Retirada, Em Andamento, Concluida e Cancelada. Transicoes invalidas sao bloqueadas pelo sistema, garantindo integridade operacional.')

    add_heading_styled(doc, '2.6 Regras de Disponibilidade', 2)
    add_body(doc, 'O administrador pode configurar regras que definem quando e onde as viagens podem ser solicitadas. Cada regra combina: area geografica (coordenada central e raio em km), dias da semana permitidos, e faixa horaria. O sistema calcula a distancia entre o ponto de partida e o centro da area permitida para validar a solicitacao. Multiplas regras podem coexistir simultaneamente.')

    add_heading_styled(doc, '2.7 Check-in e Check-out de Veiculos', 2)
    add_body(doc, 'Como os veiculos pertencem a empresa, o sistema implementa um controle formal de retirada e devolucao. O motorista registra a retirada do veiculo no inicio do turno (com quilometragem e nivel de combustivel) e a devolucao ao final. O sistema impede que um veiculo ja retirado por outro motorista seja alocado, evitando conflitos operacionais.')

    add_heading_styled(doc, '2.8 Rastreamento em Tempo Real', 2)
    add_body(doc, 'Um servico dedicado de WebSocket fornece rastreamento em tempo real da posicao dos veiculos. O dashboard administrativo pode receber atualizacoes de toda a frota simultaneamente, enquanto o passageiro pode acompanhar a posicao do veiculo durante sua viagem. O servico esta preparado para integracao com APIs de rastreadores GPS comerciais.')

    add_heading_styled(doc, '2.9 Relatorios Operacionais', 2)
    add_body(doc, 'Relatorios de viagens podem ser exportados nos formatos Excel (XLSX), CSV e JSON, com filtros por periodo, status e centro de custo. Os relatorios incluem dados completos de origem, destino, passageiro, motorista, veiculo, centro de custo, horarios e duracao das viagens.')

    add_heading_styled(doc, '2.10 Auditoria e Compliance', 2)
    add_body(doc, 'Todas as acoes de criacao, alteracao e exclusao em dados sensiveis sao registradas em trilha de auditoria imutavel (append-only). Cada registro inclui: usuario que executou a acao, tipo de acao, recurso afetado, dados alterados, endereco IP e navegador utilizado. Os logs podem ser consultados com filtros avancados e paginacao, restritos ao administrador global, garantindo compliance empresarial.')

    # 3
    add_heading_styled(doc, '3. Arquitetura da Solucao', 1)
    add_body(doc, 'A plataforma utiliza uma arquitetura moderna baseada em Next.js, com tipagem estatica (TypeScript), banco de dados relacional via Prisma ORM e interface responsiva com Tailwind CSS e componentes shadcn/ui.')
    add_table(doc, ['Camada', 'Tecnologia'], [
        ['Framework Web', 'Next.js 16 (App Router)'],
        ['Linguagem', 'TypeScript (tipagem estatica)'],
        ['Banco de Dados', 'Prisma ORM (compativel com PostgreSQL, MySQL, SQLite)'],
        ['Interface', 'Tailwind CSS + shadcn/ui'],
        ['Autenticacao', 'JWT (JSON Web Tokens) com refresh'],
        ['Tempo Real', 'WebSocket (Socket.io)'],
        ['Exportacao', 'SheetJS (XLSX/CSV)'],
    ])
    add_body(doc, 'O banco de dados esta pronto para migracao para PostgreSQL em producao sem alteracoes de codigo, bastando alterar a string de conexao.')

    # 4
    add_heading_styled(doc, '4. Perfis de Acesso', 1)
    add_table(doc, ['Perfil', 'Descricao', 'Acesso'], [
        ['Super Administrador', 'Administracao global', 'Acesso total a todas as funcionalidades, incluindo configuracoes e auditoria.'],
        ['Gerente', 'Gestao de filial ou unidade', 'Gerencia usuarios, veiculos, viagens e relatorios da sua unidade.'],
        ['Motorista', 'Operador de viagens', 'Visualiza e atualiza o status das viagens a ele atribuidas.'],
        ['Passageiro', 'Solicitante de viagens', 'Solicita viagens e acompanha suas solicitacoes.'],
    ])

    # 5
    add_heading_styled(doc, '5. Deploy e Infraestrutura', 1)
    add_body(doc, 'A aplicacao e configurada para deploy automatizado na Vercel a partir do repositorio GitHub. Basta conectar o repositorio, configurar as variaveis de ambiente (banco de dados e chaves de autenticacao), e a plataforma cuida do processo de build e deploy automaticamente.')
    add_body(doc, 'Para producao, recomenda-se utilizar um banco de dados gerenciado (PostgreSQL) e configurar chaves criptograficas fortes para os tokens de autenticacao. O servico de rastreamento em tempo real pode ser deployado separadamente em plataformas cloud.')

    # 6
    add_heading_styled(doc, '6. Entregaveis', 1)
    add_table(doc, ['Item', 'Descricao'], [
        ['Codigo-fonte', 'Repositorio GitHub com backend, dashboard e servico WebSocket.'],
        ['Database Schema', 'Schema com 12 entidades, relacoes e indexes.'],
        ['API REST', '26 endpoints com autenticacao, validacao e auditoria.'],
        ['Dashboard', 'Interface administrativa web responsiva com 11 modulos.'],
        ['Servico WebSocket', 'Microservico de rastreamento em tempo real.'],
        ['Dados de Demo', 'Script de populacao com dados de demonstracao.'],
    ])

    # 7
    add_heading_styled(doc, '7. Proximos Passos Sugeridos', 1)
    add_body(doc, 'Integracao com sistema de rastreamento GPS corporativo existente. Desenvolvimento de aplicativo mobile para motoristas e passageiros. Migracao para banco de dados PostgreSQL em nuvem. Implementacao de notificacoes push e email. Configuracao de ambiente de homologacao. Documentacao complementar da API no formato OpenAPI. Testes automatizados de integracao.')

    path = os.path.join(OUTPUT_DIR, 'FleetControl-Entrega-Cliente.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    create_dev_doc()
    create_client_doc()
    print('\nDone!')
